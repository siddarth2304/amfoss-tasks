import os
import csv
import docx
from googleapiclient.discovery import build
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes, ConversationHandler
)
import logging

# Setup logging
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

# Google Books API setup
API_KEY = 'AIzaSyCirb33eHOe_n92OxK9HEtSgU-jE8ihfv8'  # Replace with your actual Google Books API key
service = build('books', 'v1', developerKey=API_KEY)

reading_list = []

# Conversation states
BOOK_GENRE, BOOK_PREVIEW, ADD_TO_LIST, ADD_BOOK_NAME = range(4)

# /start command
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text('Welcome to PagePal! Use /help to see available commands.')

# /help command
async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    help_text = (
        "/start - Start the bot\n"
        "/book - Get book recommendations by genre\n"
        "/preview - Get a preview link for a book\n"
        "/list - Manage your reading list\n"
        "/reading_list - View your reading list\n"
    )
    await update.message.reply_text(help_text)

# /book command
async def book(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text('Please enter the genre you are interested in:')
    return BOOK_GENRE

async def genre_input(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    genre = update.message.text
    results = service.volumes().list(q=f'subject:{genre}', maxResults=5).execute()
    books = []

    for item in results.get('items', []):
        title = item['volumeInfo'].get('title', 'N/A')
        authors = ', '.join(item['volumeInfo'].get('authors', ['N/A']))
        description = item['volumeInfo'].get('description', 'N/A')
        published_date = item['volumeInfo'].get('publishedDate', 'N/A')
        language = item['volumeInfo'].get('language', 'N/A')
        preview_link = item['volumeInfo'].get('previewLink', 'N/A')
        
        books.append([title, authors, description, published_date, language, preview_link])
    
    with open('books.csv', 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(['Title', 'Author', 'Description', 'Published Date', 'Language', 'Preview Link'])
        writer.writerows(books)

    await update.message.reply_document(document=open('books.csv', 'rb'))
    return ConversationHandler.END

# /preview command
async def preview(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text('Please enter the book name for preview:')
    return BOOK_PREVIEW

async def book_preview(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    book_name = update.message.text
    results = service.volumes().list(q=book_name, maxResults=1).execute()

    if results.get('items', []):
        preview_link = results['items'][0]['volumeInfo'].get('previewLink', 'N/A')
        await update.message.reply_text(f'Preview link: {preview_link}')
    else:
        await update.message.reply_text('Sorry, no preview available for this book.')
    return ConversationHandler.END

# /list command
async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text('Please enter the book name to add to your reading list:')
    return ADD_TO_LIST

async def add_to_list(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    book_name = update.message.text
    results = service.volumes().list(q=book_name, maxResults=1).execute()

    if results.get('items', []):
        title = results['items'][0]['volumeInfo'].get('title', 'N/A')
        preview_link = results['items'][0]['volumeInfo'].get('previewLink', 'N/A')
        reading_list.append((title, preview_link))
        await update.message.reply_text(f'Added {title} to your reading list!')
    else:
        await update.message.reply_text('Sorry, no information available for this book.')
    return ConversationHandler.END

# /reading_list command
async def reading_list_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    buttons = [
        [InlineKeyboardButton('Add a book', callback_data='add')],
        [InlineKeyboardButton('Delete a book', callback_data='delete')],
        [InlineKeyboardButton('View Reading List', callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(buttons)
    await update.message.reply_text('Choose an option:', reply_markup=reply_markup)
    return ADD_BOOK_NAME

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()

    if query.data == 'add':
        await query.message.reply_text('Please enter the book name to add:')
        return ADD_BOOK_NAME

    elif query.data == 'view':
        doc = docx.Document()
        doc.add_heading('Reading List', 0)

        for book in reading_list:
            doc.add_paragraph(f"Title: {book[0]}\nPreview Link: {book[1]}")

        doc.save('reading_list.docx')
        await query.message.reply_document(document=open('reading_list.docx', 'rb'))

    elif query.data == 'delete':
        # Implement delete functionality
        if reading_list:
            last_book = reading_list.pop()  # Remove the last added book
            await query.message.reply_text(f'Removed {last_book[0]} from your reading list.')
        else:
            await query.message.reply_text('Your reading list is empty.')
    
    return ConversationHandler.END

async def add_book_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    book_name = update.message.text
    results = service.volumes().list(q=book_name, maxResults=1).execute()

    if results.get('items', []):
        title = results['items'][0]['volumeInfo'].get('title', 'N/A')
        preview_link = results['items'][0]['volumeInfo'].get('previewLink', 'N/A')
        reading_list.append((title, preview_link))
        await update.message.reply_text(f'Added {title} to your reading list!')
    else:
        await update.message.reply_text('Sorry, no information available for this book.')

    return ConversationHandler.END

# Main function to start the bot
def run_bot():
    application = Application.builder().token("7407411246:AAHVZF8inADb1tt0ka_Mx19unp2Y9Ca6AYI").build()

    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('book', book),
            CommandHandler('preview', preview),
            CommandHandler('list', list_command),
            CommandHandler('reading_list', reading_list_command),
        ],
        states={
            BOOK_GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, genre_input)],
            BOOK_PREVIEW: [MessageHandler(filters.TEXT & ~filters.COMMAND, book_preview)],
            ADD_TO_LIST: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_to_list)],
            ADD_BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_book_name)],
        },
        fallbacks=[CommandHandler('start', start)],
    )

    application.add_handler(conv_handler)
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("reading_list", reading_list_command))
    application.add_handler(CallbackQueryHandler(button_handler))

    # Start polling and handle incoming updates
    application.run_polling()

if __name__ == '__main__':
    run_bot()
